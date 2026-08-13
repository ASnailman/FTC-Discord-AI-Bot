"""Turns validated `ingest.IngestedFile`s into plain text and/or images.

Per-format extraction, each bounded so a hostile or oversized file can't
blow up memory or the token budget:

- PDF: text via `pypdf` first. If a PDF has no extractable text at all --
  the reference world-championship portfolio this feature is styled on is
  exactly this shape, an 85-page image-only export -- a bounded number of
  pages are rasterized with `pypdfium2` instead, so the content still
  reaches the model as images for a vision pass (see vision.py) rather
  than silently producing nothing.
- .docx: paragraph text via `python-docx`.
- .md/.txt: decoded as UTF-8 (`errors="replace"` -- never raises on bad
  bytes).
- Images (.png/.jpg/.jpeg/.webp/.gif): decoded and normalized (converted
  to RGB, EXIF stripped by re-encoding, downscaled to
  `config.PORTFOLIO_IMAGE_MAX_EDGE_PX`) for embedding and vision analysis.

`Image.MAX_IMAGE_PIXELS` is capped so a decompression-bomb image (a small
file that decodes to an enormous bitmap) is rejected by Pillow itself
rather than exhausting memory.
"""
import io
from dataclasses import dataclass, field

import pypdf
import pypdfium2 as pdfium
from docx import Document
from PIL import Image

import config
from logging_setup import get_logger
from .ingest import IngestedFile

logger = get_logger(__name__)

Image.MAX_IMAGE_PIXELS = 40_000_000


@dataclass(frozen=True)
class ExtractedText:
    filename: str
    text: str


@dataclass(frozen=True)
class ExtractedImage:
    filename: str
    image: "Image.Image"
    source: str  # "upload" | "pdf_page"


@dataclass(frozen=True)
class ExtractionResult:
    texts: list[ExtractedText] = field(default_factory=list)
    images: list[ExtractedImage] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _normalize_image(img: "Image.Image") -> "Image.Image":
    img = img.convert("RGB")
    max_edge = config.PORTFOLIO_IMAGE_MAX_EDGE_PX
    if max(img.size) > max_edge:
        ratio = max_edge / max(img.size)
        img = img.resize((max(1, round(img.width * ratio)), max(1, round(img.height * ratio))))
    return img


def _cap_text(text: str) -> str:
    return text[: config.PORTFOLIO_MAX_EXTRACTED_CHARS_PER_FILE]


def _extract_pdf(file: IngestedFile) -> ExtractionResult:
    warnings: list[str] = []
    reader = pypdf.PdfReader(io.BytesIO(file.data))
    num_pages = len(reader.pages)
    pages_to_read = min(num_pages, config.PORTFOLIO_MAX_PDF_PAGES)
    if num_pages > pages_to_read:
        warnings.append(
            f"{file.filename}: {num_pages} pages exceeds the {config.PORTFOLIO_MAX_PDF_PAGES}-page "
            f"cap -- only the first {pages_to_read} were read."
        )

    text_parts = []
    for i in range(pages_to_read):
        try:
            text_parts.append(reader.pages[i].extract_text() or "")
        except Exception:
            logger.warning("failed to extract text from %s page %d", file.filename, i, exc_info=True)

    text = "\n".join(part for part in text_parts if part.strip())
    images: list[ExtractedImage] = []

    if not text.strip():
        render_n = min(num_pages, config.PORTFOLIO_MAX_PDF_RENDER_PAGES)
        warnings.append(
            f"{file.filename}: no extractable text (image-only PDF export) -- "
            f"rendering the first {render_n} page(s) for visual analysis instead."
        )
        try:
            pdf = pdfium.PdfDocument(file.data)
            for i in range(render_n):
                page = pdf[i]
                bitmap = page.render(scale=150 / 72)
                pil_image = _normalize_image(bitmap.to_pil())
                images.append(ExtractedImage(filename=f"{file.filename}#page{i + 1}", image=pil_image, source="pdf_page"))
        except Exception:
            logger.warning("failed to rasterize %s", file.filename, exc_info=True)
            warnings.append(f"{file.filename}: could not be rendered as images either -- skipped.")

    texts = [ExtractedText(filename=file.filename, text=_cap_text(text))] if text.strip() else []
    return ExtractionResult(texts=texts, images=images, warnings=warnings)


def _extract_docx(file: IngestedFile) -> ExtractionResult:
    document = Document(io.BytesIO(file.data))
    text = "\n".join(p.text for p in document.paragraphs)
    texts = [ExtractedText(filename=file.filename, text=_cap_text(text))] if text.strip() else []
    return ExtractionResult(texts=texts)


def _extract_plain_text(file: IngestedFile) -> ExtractionResult:
    text = file.data.decode("utf-8", errors="replace")
    texts = [ExtractedText(filename=file.filename, text=_cap_text(text))] if text.strip() else []
    return ExtractionResult(texts=texts)


def _extract_image(file: IngestedFile) -> ExtractionResult:
    img = _normalize_image(Image.open(io.BytesIO(file.data)))
    return ExtractionResult(images=[ExtractedImage(filename=file.filename, image=img, source="upload")])


_EXTRACTORS = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "md": _extract_plain_text,
    "txt": _extract_plain_text,
    "png": _extract_image,
    "jpg": _extract_image,
    "jpeg": _extract_image,
    "webp": _extract_image,
    "gif": _extract_image,
}


def extract_all(files: list[IngestedFile]) -> ExtractionResult:
    texts: list[ExtractedText] = []
    images: list[ExtractedImage] = []
    warnings: list[str] = []
    total_chars = 0

    for file in files:
        extractor = _EXTRACTORS.get(file.extension)
        if extractor is None:
            warnings.append(f"{file.filename}: unsupported file type -- skipped.")
            continue
        try:
            result = extractor(file)
        except Exception:
            logger.warning("failed to extract %s", file.filename, exc_info=True)
            warnings.append(f"{file.filename}: could not be read (corrupt or unreadable) -- skipped.")
            continue

        for t in result.texts:
            remaining = config.PORTFOLIO_MAX_EXTRACTED_CHARS_TOTAL - total_chars
            if remaining <= 0:
                warnings.append("Total extracted-text budget reached -- remaining files' text was skipped.")
                break
            trimmed = t.text[:remaining]
            total_chars += len(trimmed)
            texts.append(ExtractedText(filename=t.filename, text=trimmed))

        images.extend(result.images)
        warnings.extend(result.warnings)

    if len(images) > config.PORTFOLIO_MAX_VISION_IMAGES:
        warnings.append(f"Only the first {config.PORTFOLIO_MAX_VISION_IMAGES} images will be analyzed/embedded.")
        images = images[: config.PORTFOLIO_MAX_VISION_IMAGES]

    return ExtractionResult(texts=texts, images=images, warnings=warnings)
